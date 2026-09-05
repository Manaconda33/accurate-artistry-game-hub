from pathlib import Path
import importlib.util
import re

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
BASE_SCRIPT = ROOT / "tools/slice5_docs_sync.py"

spec = importlib.util.spec_from_file_location("slice5_docs_sync", BASE_SCRIPT)
if spec is None or spec.loader is None:
    raise RuntimeError("Could not load Slice 5 sync module")
module = importlib.util.module_from_spec(spec)
spec.loader.exec_module(module)


def patch_docx() -> None:
    docx_path = ROOT / "docs/Manacondas_Minigame_Mayhem_PRD_v1.1.docx"
    doc = Document(docx_path)

    # Preserve the current cover/version format. If the DOCX already carries a
    # working-amendment phrase, advance it; otherwise the inserted approved
    # amendment heading is the authoritative visible revision marker.
    version_pattern = re.compile(r"working implementation amendment\s+\d+(?:\.\d+)?", re.I)
    for paragraph in doc.paragraphs:
        if version_pattern.search(paragraph.text):
            paragraph.text = version_pattern.sub(
                "working implementation amendment 2.2",
                paragraph.text,
                count=1,
            )
            break

    if not any(module.AMENDMENT_HEADING in paragraph.text for paragraph in doc.paragraphs):
        contents = next(
            (paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "Contents"),
            None,
        )
        reference_heading = next(
            (
                paragraph
                for paragraph in doc.paragraphs
                if "Approved implementation amendment 2.1" in paragraph.text
            ),
            None,
        )
        if contents is None or reference_heading is None:
            anchors = [
                paragraph.text
                for paragraph in doc.paragraphs
                if "amendment" in paragraph.text.lower() or paragraph.text.strip() == "Contents"
            ]
            raise RuntimeError(f"DOCX amendment insertion anchors not found: {anchors[:12]}")

        heading_style = reference_heading.style.name if reference_heading.style else None
        style_names = {style.name for style in doc.styles}
        normal_style = "Normal" if "Normal" in style_names else None
        contents.insert_paragraph_before(module.AMENDMENT_HEADING, style=heading_style)
        for block in module.DOCX_BLOCKS:
            contents.insert_paragraph_before(block, style=normal_style)

    doc.save(docx_path)


def main() -> None:
    module.patch_markdown()
    patch_docx()
    for temporary in [
        ROOT / ".github/workflows/slice5-docs-sync.yml",
        ROOT / "tools/slice5_docs_sync.py",
        ROOT / "tools/slice5_docs_sync_fix.py",
        ROOT / "tools/slice5_docs_sync_fix2.py",
    ]:
        if temporary.exists():
            temporary.unlink()


if __name__ == "__main__":
    main()
